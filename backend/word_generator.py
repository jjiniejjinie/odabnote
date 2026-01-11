"""
Word 문서 생성 모듈
문제지 Word, 정답지 Word 생성
완벽한 2x3 테이블 레이아웃
"""
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import requests
from io import BytesIO
import cloudinary


class WordGenerator:
    """Word 문서 생성기 클래스"""
    
    def __init__(self):
        """Word 생성기 초기화"""
        pass
    
    def _download_image(self, image_path):
        """Cloudinary에서 이미지 다운로드"""
        try:
            image_url = cloudinary.CloudinaryImage(image_path).build_url()
            response = requests.get(image_url, timeout=10)
            if response.status_code == 200:
                return BytesIO(response.content)
            return None
        except Exception as e:
            print(f"Image download error: {e}")
            return None
    
    def _add_watermark_footer(self, doc, username=None):
        """푸터에 워터마크 추가"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        
        watermark_text = "개인 학습용"
        if username:
            watermark_text += f" - {username}"
        watermark_text += " - 재배포 금지"
        
        footer_para.text = watermark_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 스타일 설정
        run = footer_para.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _set_cell_border(self, cell, **kwargs):
        """셀 테두리 설정"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # 테두리 생성
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '12')  # 1pt
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
        
        tcPr.append(tcBorders)
    
    def generate_problem_docx(self, unit, problems, output_path, username=None):
        """
        문제지 Word 생성 (완벽한 2x3 레이아웃)
        
        Args:
            unit: Unit 모델 객체
            problems: Problem 모델 객체 리스트
            output_path: 저장할 파일 경로
            username: 사용자명 (워터마크용, 선택사항)
        """
        doc = Document()
        
        # 페이지 설정
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        
        # 제목과 날짜
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{unit.workbook.name} - {unit.name}")
        title_run.font.size = Pt(16)
        title_run.bold = True
        
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        date_run.font.size = Pt(10)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # 여백
        
        # 3문제씩 묶어서 처리
        problems_per_page = 3
        
        for page_idx in range(0, len(problems), problems_per_page):
            page_problems = problems[page_idx:page_idx + problems_per_page]
            
            # 2x3 표 생성 (행: 문제 수, 열: 2)
            table = doc.add_table(rows=len(page_problems), cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # 표 전체 너비 설정
            table.width = Cm(19)
            
            for idx, problem in enumerate(page_problems):
                row = table.rows[idx]
                
                # 행 높이 고정
                row.height = Cm(9)
                
                # 왼쪽 셀 (문제)
                left_cell = row.cells[0]
                left_cell.width = Cm(9.5)
                self._set_cell_border(left_cell)
                
                # 문제 번호
                left_para = left_cell.paragraphs[0]
                left_run = left_para.add_run(f"문제 {problem.problem_number}")
                left_run.bold = True
                left_run.font.size = Pt(12)
                
                # 이미지 추가
                if problem.problem_image_path:
                    img_data = self._download_image(problem.problem_image_path)
                    if img_data:
                        try:
                            left_cell.add_paragraph()  # 공백
                            img_para = left_cell.add_paragraph()
                            img_run = img_para.add_run()
                            img_run.add_picture(img_data, width=Cm(8.5))
                        except Exception as e:
                            error_para = left_cell.add_paragraph()
                            error_para.add_run(f"[이미지 로드 실패]")
                    else:
                        error_para = left_cell.add_paragraph()
                        error_para.add_run(f"[이미지 없음]")
                
                # 오른쪽 셀 (풀이 공간)
                right_cell = row.cells[1]
                right_cell.width = Cm(9.5)
                self._set_cell_border(right_cell)
                
                right_para = right_cell.paragraphs[0]
                right_run = right_para.add_run("풀이")
                right_run.bold = True
                right_run.font.size = Pt(11)
                
                # 풀이 공간 (여러 줄)
                for _ in range(10):
                    right_cell.add_paragraph()
            
            # 페이지 나누기 (마지막 페이지 제외)
            if page_idx + problems_per_page < len(problems):
                doc.add_page_break()
        
        # 워터마크 추가
        self._add_watermark_footer(doc, username)
        
        # 저장
        doc.save(output_path)
        return output_path
    
    def generate_answer_docx(self, unit, problems, output_path, username=None):
        """
        정답지 Word 생성 (2열 레이아웃)
        
        Args:
            unit: Unit 모델 객체
            problems: Problem 모델 객체 리스트
            output_path: 저장할 파일 경로
            username: 사용자명 (워터마크용, 선택사항)
        """
        doc = Document()
        
        # 페이지 설정
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        
        # 제목
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{unit.workbook.name} - {unit.name} [정답]")
        title_run.font.size = Pt(16)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 여백
        
        # 정답이 있는 문제만 필터링
        problems_with_answers = [p for p in problems if p.has_answer]
        
        if not problems_with_answers:
            no_answer_para = doc.add_paragraph()
            no_answer_run = no_answer_para.add_run("정답이 등록된 문제가 없습니다.")
            no_answer_run.font.size = Pt(12)
            no_answer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 2개씩 묶어서 2열 테이블로 배치
            for i in range(0, len(problems_with_answers), 2):
                # 테이블 생성 (1행 2열)
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.width = Cm(18)
                
                row = table.rows[0]
                
                # 왼쪽 문제
                if i < len(problems_with_answers):
                    problem = problems_with_answers[i]
                    left_cell = row.cells[0]
                    left_cell.width = Cm(9)
                    self._set_cell_border(left_cell)
                    
                    self._add_answer_to_cell(left_cell, problem)
                
                # 오른쪽 문제
                if i + 1 < len(problems_with_answers):
                    problem = problems_with_answers[i + 1]
                    right_cell = row.cells[1]
                    right_cell.width = Cm(9)
                    self._set_cell_border(right_cell)
                    
                    self._add_answer_to_cell(right_cell, problem)
                
                doc.add_paragraph()  # 표 사이 여백
        
        # 워터마크 추가
        self._add_watermark_footer(doc, username)
        
        # 저장
        doc.save(output_path)
        return output_path
    
    def _add_answer_to_cell(self, cell, problem):
        """셀에 정답 추가"""
        # 문제 번호
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"문제 {problem.problem_number} 정답")
        title_run.bold = True
        title_run.font.size = Pt(11)
        
        cell.add_paragraph()  # 여백
        
        # 정답 이미지
        if problem.answer_image_path:
            img_data = self._download_image(problem.answer_image_path)
            if img_data:
                try:
                    img_para = cell.add_paragraph()
                    img_run = img_para.add_run()
                    img_run.add_picture(img_data, width=Cm(8))
                except Exception as e:
                    error_para = cell.add_paragraph()
                    error_para.add_run(f"[이미지 로드 실패]")
            else:
                error_para = cell.add_paragraph()
                error_para.add_run(f"[이미지 없음]")
        
        # 정답 텍스트
        if problem.answer_text:
            cell.add_paragraph()  # 여백
            text_para = cell.add_paragraph()
            text_para.add_run(problem.answer_text)
